"""Generate the FlexRadio Scheduler transfer instructions as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level, size in [('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12)]:
    s = doc.styles[level]
    s.font.name = 'Arial'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# Code style
code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(10)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Inches(0.5)

# --- Helper ---
def add_code(text):
    doc.add_paragraph(text, style='Code')

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    p.add_run(text)

# ============================================================
# CONTENT
# ============================================================

# Title
title = doc.add_heading('FlexRadio Scheduler', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Transfer & Installation Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# --- Prerequisites ---
doc.add_heading('Prerequisites', level=1)

doc.add_heading('Install Python', level=2)
p = doc.add_paragraph('Download and install ')
run = p.add_run('Python 3.10 or newer')
run.bold = True
p.add_run(' from python.org/downloads.')
doc.add_paragraph()
add_bullet('During installation, check the box labeled "Add Python to PATH".')
add_bullet('Verify installation by opening a terminal and running:')
add_code('python --version')

# --- Step 1: Copy ---
doc.add_heading('Step 1: Copy the Project Folder', level=1)
doc.add_paragraph(
    'Copy the entire "FlexRadio Scheduler" folder to the new computer. '
    'You can use a USB drive, network share, or zip the folder.'
)

doc.add_heading('Required Files', level=2)
doc.add_paragraph('The following files must be included:')
add_code(
    'FlexRadio Scheduler/\n'
    '    app.py\n'
    '    models.py\n'
    '    radio_monitor.py\n'
    '    flexradio.py\n'
    '    relay_agent.py\n'
    '    zerotier.py\n'
    '    requirements.txt\n'
    '    static/\n'
    '        style.css\n'
    '    templates/\n'
    '        base.html\n'
    '        login.html\n'
    '        register.html\n'
    '        profile.html\n'
    '        schedule.html\n'
    '        admin.html\n'
    '        admin_radio.html'
)

doc.add_heading('Optional: Transfer Your Data', level=2)
add_bullet(
    ' - Copy this file to bring over all users, bookings, radio config, '
    'and blocked slots to the new computer.',
    bold_prefix='instance/scheduler.db'
)
add_bullet(
    'If you do not copy the database, the application will create a fresh '
    'database on first run. You will need to re-register users and reconfigure settings.'
)

# --- Step 2: Install deps ---
doc.add_heading('Step 2: Install Dependencies', level=1)
doc.add_paragraph(
    'Open a terminal (Command Prompt or PowerShell) and navigate to the project folder:'
)
add_code('cd "C:\\path\\to\\FlexRadio Scheduler"')
doc.add_paragraph('Then install the required Python packages:')
add_code('pip install -r requirements.txt')
doc.add_paragraph('This installs Flask, Flask-SQLAlchemy, Flask-Login, and the requests library.')

# --- Step 3: Start ---
doc.add_heading('Step 3: Start the Application', level=1)
doc.add_paragraph('From the project folder, run:')
add_code('python app.py')
doc.add_paragraph('You should see output similar to:')
add_code(
    '* Serving Flask app \'app\'\n'
    '* Debug mode: on\n'
    '* Running on http://127.0.0.1:5000'
)
doc.add_paragraph(
    'Open a web browser and navigate to http://127.0.0.1:5000 to access the scheduler.'
)

# --- Step 4: Configure ---
doc.add_heading('Step 4: Initial Configuration', level=1)

doc.add_heading('If You Copied the Database', level=2)
doc.add_paragraph(
    'All your settings, users, and bookings carry over automatically. '
    'However, you may need to update the Radio IP Address if the new computer '
    'is on a different network.'
)
add_bullet('Log in with your existing admin account.')
add_bullet('Go to Admin > Radio Configuration.')
add_bullet('Verify or update the Radio IP Address for the new network.')

doc.add_heading('If Starting Fresh (No Database)', level=2)
add_bullet('Register a new account. The first user automatically becomes the admin.')
add_bullet('Navigate to Admin > Radio Configuration to set up:')

p = doc.add_paragraph(style='List Bullet 2')
p.add_run('Connection mode').bold = True
p.add_run(' (Direct LAN or Remote Relay)')
p = doc.add_paragraph(style='List Bullet 2')
p.add_run('Radio IP Address').bold = True
p.add_run(' and TCP port')
p = doc.add_paragraph(style='List Bullet 2')
p.add_run('Schedule enforcement').bold = True
p.add_run(' toggle')
p = doc.add_paragraph(style='List Bullet 2')
p.add_run('ZeroTier').bold = True
p.add_run(' network access control (if used)')
p = doc.add_paragraph(style='List Bullet 2')
p.add_run('Monthly hour limits').bold = True
p.add_run(' and check interval')

# --- LAN Access ---
doc.add_heading('Step 5: Allow LAN Access (Optional)', level=1)
doc.add_paragraph(
    'By default, the server only listens on localhost (127.0.0.1). '
    'To make it accessible from other devices on your local network:'
)

doc.add_heading('Option A: Change the Bind Address', level=2)
doc.add_paragraph(
    'Edit the last line of app.py and change it to:'
)
add_code("app.run(debug=True, host='0.0.0.0', port=5000)")
doc.add_paragraph(
    'This binds the server to all network interfaces. Other devices can then '
    'access the scheduler at http://<your-computer-ip>:5000.'
)

doc.add_heading('Option B: Use a Production Server', level=2)
doc.add_paragraph('For a more robust setup, use Waitress (Windows-compatible):')
add_code('pip install waitress')
add_code('waitress-serve --host=0.0.0.0 --port=5000 app:app')

add_note(
    'You may need to allow port 5000 through your Windows Firewall. '
    'Go to Windows Defender Firewall > Advanced Settings > Inbound Rules > New Rule > Port > TCP 5000.'
)

# --- Relay Agent ---
doc.add_heading('Step 6: Relay Agent Setup (If Used)', level=1)
doc.add_paragraph(
    'If you are using the Remote Relay mode (scheduler is on a different network '
    'than the radio), you need to run the relay agent on a computer at the radio site.'
)

add_bullet('Copy relay_agent.py to the computer at the radio site.')
add_bullet('Install the requirements: pip install flask requests')
add_bullet('Start the relay agent:')
add_code('python relay_agent.py --radio-ip 192.168.1.100 --api-key YOUR_SECRET_KEY')

doc.add_paragraph('Relay agent options:')
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.columns[0].width = Inches(2)
table.columns[1].width = Inches(4.5)

headers = table.rows[0].cells
headers[0].text = 'Option'
headers[1].text = 'Description'
for cell in headers:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ('--radio-ip', 'FlexRadio IP on the LAN (required)'),
    ('--api-key', 'Shared secret key (required, must match scheduler config)'),
    ('--radio-port', 'Radio TCP port (default: 4992)'),
    ('--host', 'Relay bind address (default: 0.0.0.0)'),
    ('--port', 'Relay HTTP port (default: 5001)'),
]
for i, (opt, desc) in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = opt
    row[1].text = desc
    # Make option name monospace
    for paragraph in row[0].paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)

# --- Troubleshooting ---
doc.add_heading('Troubleshooting', level=1)

doc.add_heading('Common Issues', level=2)

p = doc.add_paragraph()
run = p.add_run('"pip" is not recognized')
run.bold = True
doc.add_paragraph(
    'Python was installed without adding to PATH. Reinstall Python and check '
    '"Add Python to PATH", or use the full path: '
    'C:\\Users\\<you>\\AppData\\Local\\Programs\\Python\\Python3xx\\Scripts\\pip.exe'
)

p = doc.add_paragraph()
run = p.add_run('Port 5000 already in use')
run.bold = True
doc.add_paragraph(
    'Another application is using port 5000. Either stop that application or '
    'change the port in app.py: app.run(debug=True, port=5001)'
)

p = doc.add_paragraph()
run = p.add_run('Cannot connect to radio')
run.bold = True
doc.add_paragraph(
    'Verify the Radio IP Address in Admin > Radio Configuration. '
    'Ensure the scheduler computer can reach the radio on port 4992 (or relay on port 5001). '
    'Check firewalls on both machines.'
)

p = doc.add_paragraph()
run = p.add_run('Database errors after transfer')
run.bold = True
doc.add_paragraph(
    'If you see migration errors, delete the instance/scheduler.db file and restart. '
    'The application will create a fresh database with all the latest schema changes.'
)

# --- Save ---
output_path = r'C:\Users\jmajors.HINDS\Documents\Programs\FlexRadio Scheduler\FlexRadio_Scheduler_Transfer_Guide.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
